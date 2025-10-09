import sys
import os
from pathlib import Path
import re
from datetime import datetime
from git import Repo
from docx import Document
from spire.doc import *
from spire.doc.common import *

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_DIR = SCRIPT_DIR.parent
DOCS_DIR = REPO_DIR / "docs"
CSS_URL = "https://wdpprepository.org/static/css/project.css"
REMOTE = "origin"
BRANCH = "main"

def _debug_rewrite_image_paths(html: str, media_dir: Path, output_stem: str, debug: bool = True) -> str:
    media_abs = str(media_dir.resolve())
    media_win = media_abs.replace('/', '\\')
    media_posix = media_abs.replace('\\', '/')
    pat = re.compile(rf"{re.escape(media_win)}[\\/]|{re.escape(media_posix)}/", flags=re.IGNORECASE)
    new_html, n_subs = pat.subn(f"{output_stem}_files/", html)
    if debug:
        print(f"[DEBUG] Replaced {n_subs} image paths.")
    return new_html

def get_title_from_word(docx_path: Path) -> str:
    try:
        doc = Document(docx_path)
        core_props = doc.core_properties
        title_prop = (core_props.title or "").strip()

        print(f"🔍 Checking title property in {docx_path.name}...")
        if title_prop:
            print(f"✅ Found title property: '{title_prop}'")
            return title_prop

        print("⚠️  No title property found, checking for Heading 1...")
        for p in doc.paragraphs:
            style_name = getattr(p.style, "name", "").lower()
            if style_name.startswith("heading 1"):
                text = p.text.strip()
                if text:
                    print(f"✅ Found Heading 1: '{text}'")
                    return text

        print("⚠️  No Heading 1 found, using filename stem.")
        return docx_path.stem

    except Exception as e:
        print(f"❌ Could not extract title from {docx_path.name}: {e}")
        return docx_path.stem

def convert_docx_to_html_spire(docx_path: Path, output_path: Path, title: str):
    print(f"Converting {docx_path.name} → HTML using Spire.Doc …")
    media_dir = output_path.parent / f"{output_path.stem}_files"
    media_dir.mkdir(exist_ok=True)

    # Load and convert using Spire.Doc
    document = Document()
    document.LoadFromFile(str(docx_path))
    # Control whether to include headers and footers in the exported HTML
    document.HtmlExportOptions.HasHeadersFooters = False

    # Specify the name of the CSS file to use for styling the exported HTML
    document.HtmlExportOptions.CssStyleSheetFileName = "sample.css"

    # Set the CSS stylesheet type to external, so the HTML file links to the specified CSS file instead of embedding styles inline
    document.HtmlExportOptions.CssStyleSheetType = CssStyleSheetType.External

    # Configure image export: do not embed images inside HTML, save them to a separate folder
    document.HtmlExportOptions.ImageEmbedded = False
    document.HtmlExportOptions.ImagesPath = "Images/"

    # Export form fields as plain text instead of interactive form elements
    document.HtmlExportOptions.IsTextInputFormFieldAsText = True
    document.HtmlExportOptions.CssStyleSheetType = CssStyleSheetType.External
    document.SaveToFile(str(output_path), FileFormat.Html)

    # Read the output for post-processing
    html = output_path.read_text(encoding="utf-8")

    ## Remove unwanted inline <style> from Spire output if present
    #html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL | re.IGNORECASE)

    ## Inject CSS link
    #css_link = f'<link rel="stylesheet" href="{CSS_URL}">'
    #html = re.sub(r"</head>", css_link + "\n</head>", html, count=1, flags=re.IGNORECASE)

    # Fix image paths
    #html = _debug_rewrite_image_paths(html, media_dir, output_path.stem, debug=False)

    # Ensure title
    if not re.search(r"<h1\b", html, flags=re.IGNORECASE) and title:
        html = re.sub(r"<body([^>]*)>", rf"<body\1>\n<h1>{title}</h1>", html, count=1, flags=re.IGNORECASE)

    # Add front matter
    front_matter = f'---\nlayout: none\ntitle: "{title}"\n---\n'
    output_path.write_text(front_matter + html, encoding="utf-8")

    print(f"✅ HTML created: {output_path.relative_to(REPO_DIR)}")
    print(f"🖼️ Media saved in: {media_dir.relative_to(REPO_DIR)}")

def commit_and_push(repo_dir: Path, message: str):
    repo = Repo(repo_dir)
    repo.git.add(A=True)
    if repo.is_dirty():
        repo.index.commit(message)
        repo.git.push(REMOTE, BRANCH)
        print("🚀 Changes pushed to GitHub Pages.")
    else:
        print("No new changes detected.")

def main():
    if len(sys.argv) < 2:
        print("Usage: python publish_article2.py <word_file>")
        sys.exit(1)

    word_file = Path(sys.argv[1])
    if not word_file.is_absolute():
        word_file = (SCRIPT_DIR / word_file).resolve()

    if not word_file.exists():
        print(f"❌ Word file not found: {word_file}")
        sys.exit(1)

    title = get_title_from_word(word_file)
    DOCS_DIR.mkdir(exist_ok=True)
    output_html = DOCS_DIR / f"{word_file.stem}.html"

    convert_docx_to_html_spire(word_file, output_html, title)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    commit_and_push(REPO_DIR, f"Update {word_file.name} → HTML ({timestamp})")

    base_url = "https://asselapathirana.github.io/wdpprepository.org"
    print(f"\n🌐 Live page URL:\n{base_url}/{output_html.name}\n")

if __name__ == "__main__":
    main()
